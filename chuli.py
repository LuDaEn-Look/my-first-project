import os

import row
from docx import Document
from openpyxl import load_workbook, Workbook

import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

def load_reference_data(excel_path):
    """
    从Excel加载参考数据
    :param excel_path: Excel文件路径
    :return: 字典{学号: 学生信息}
    """
    ref_data = {}
    wb = load_workbook(excel_path)
    ws = wb.active

    # 假设Excel格式为：姓名 学号 指导老师 毕业设计题目 成果表现形式
    for row in ws.iter_rows(min_row=2, values_only=True):
        student_id = row[1]  # 学号作为字典键
        ref_data[student_id] = {  # 使用学号作为键存储学生信息
            'id': row[2],
            'name': row[1],
            'advisor': row[5], 
            'title': row[4]
        }
    # print(f"学号: {row[2]}, 姓名: {row[1]}, 指导老师: {row[5]}, 毕业设计题目: {row[4]}")
    print(ref_data)
    return ref_data

def check_document_consistency(student_folder, ref_data):
    """
    检查一个学生5个Word文档的一致性
    :param student_folder: 学生文件夹路径
    :param ref_data: 参考数据字典
    :return: 不一致的问题列表和学生ID
    """
    issues = []
    student_id = None  # 初始化为None
    doc_files = [f for f in os.listdir(student_folder) if f.endswith('.docx') or f.endswith('.doc')]
    # if len(doc_files) != 5:
    #     issues.append(f"文档数量不足5个，当前有{len(doc_files)}个")
    #     return issues, student_id  # 确保返回两个值

    # 检查每个文档
    for doc_file in doc_files:
        doc_path = os.path.join(student_folder, doc_file)
        try:
            doc = Document(doc_path)
            content = "\n".join([para.text for para in doc.paragraphs])

            # 检查禁止内容
            forbidden_phrases = ['以论文', '实习总结', '实习报告']
            for phrase in forbidden_phrases:
                if phrase in content:
                    issues.append(f"文档'{doc_file}'包含禁止内容: '{phrase}'")

            # 提取学生信息
            print(f"正在处理文档: {doc_file}")
            info = extract_student_info(doc)
            # 检查文档中提取的信息与参考数据是否匹配
            if info['name']:  # 只要文档中有学号信息就进行比对
                student_id = info['name'][0]  # 取第一个学号作为基准
                print(f"正在处理学号: {student_id}")
                # 与参考数据比对
                print(ref_data)
                if student_id in ref_data:
                    print(f"正在处理学号: {student_id}")
                    ref_info = ref_data[student_id]
                    # 检查学号是否与参考数据一致
                    if info['id'] and not any(str(id).strip() == str(ref_info['id']).strip() for id in info['id']):
                        issues.append(f"文档'{doc_file}'学号不一致: '{', '.join(info['id'])}' vs 参考'{ref_info['id']}'")
                    # 检查指导老师是否在参考数据中
                    if info['advisor'] and not any(advisor == ref_info['advisor'] for advisor in info['advisor']):
                        issues.append(f"文档'{doc_file}'指导老师不一致: '{info['advisor']}' vs 参考'{ref_info['advisor']}'")
                    # 检查毕业设计题目是否在参考数据中
                    if info['title'] and not any(title == ref_info['title'] for title in info['title']):
                        issues.append(f"文档'{doc_file}'毕业设计题目不一致: '{info['title']}' vs 参考'{ref_info['title']}'")
                else:
                    print(f"学号: {student_id}在参考数据中不存在")
                    issues.append(f"文档'{doc_file}'中姓名{student_id}在参考数据中不存在")

        except Exception as e:
            issues.append(f"文档'{doc_file}'无法读取，可能已损坏: {str(e)}")
            continue  # 跳过这个文档继续检查下一个

    return issues, student_id

def extract_student_info(doc):
    """
    从Word文档的表格和段落中提取学生信息
    :param doc: 可以是Document对象或字符串内容
    :return: 学生信息字典，值为列表形式
    """
    info = {
        'name': [],
        'id': [],
        'advisor': [],
        'title': [],
        'presentation': []
    }
    # # 处理字符串内容的情况
    # if isinstance(doc, str):
    #     content = doc
    #     lines = content.split('\n')
    #     print(lines)
    #     for line in lines:
    #         # 处理常见分隔符(:：)
    #         if '学生姓名' in line or '学生姓名：' in line:
    #             value = line.split('：')[-1].strip() if '：' in line else line.split(':')[-1].strip()
    #             if value: info['name'].append(value)
    #         elif '学号' in line or '学号：' in line:
    #             # print(f"学号行: {line}")
    #             value = line.split('：')[-1].strip() if '：' in line else line.split(':')[-1].strip()
    #             if value: info['id'].append(value)
    #         elif '指导老师：' in line or '指导老师' in line:
    #             value = line.split('：')[-1].strip() if '：' in line else line.split(':')[-1].strip()
    #             if value: info['advisor'].append(value)
    #         elif '毕业设计题目' in line or '毕业设计题目：' in line:
    #             value = line.split('：')[-1].strip() if '：' in line else line.split(':')[-1].strip()
    #             if value: info['title'].append(value)
    #
    #     return info

    # 处理Document对象 - 支持多列表格
    if hasattr(doc, 'tables'):
        for table in doc.tables:
            for row in table.rows:
                # 处理多列表格(偶数列)
                for i in range(0, len(row.cells)-1, 2):  # 每次步进2列
                    if i+1 >= len(row.cells):  # 确保有值列
                        continue

                    key = row.cells[i].text.strip().replace(' ', '').replace('\t', '')
                    value = row.cells[i+1].text.strip()

                    # 处理常见键名变体 - 要求全文匹配
                    if key in ['学生姓名：', '学生姓名', '名字']:
                        if value and value not in info['name']:
                            info['name'].append(value.replace(' ', ''))  # 去除姓名中的空格
                    elif key == '学号':
                        if value and value not in info['id']:
                            info['id'].append(value)
                    elif key in ['指导教师', '指导老师', '指导老师：']:
                        if value and value not in info['advisor']:
                            info['advisor'].append(value)
                    elif key in ['毕业设计题目', '毕业设计题目：']:
                        if value and value not in info['title']:
                            info['title'].append(value)
                    elif key in ['成果表现形式', '成果形式', '表现形式']:
                        if value and value not in info['presentation']:
                            info['presentation'].append(value)

    # print(info)
    return info


class TableViewerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("毕业设计文档检查工具")
        self.ref_data = None

        # 创建主框架
        self.main_frame = ttk.Frame(root, padding="10")
        self.main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        # 参考Excel文件选择
        ttk.Label(self.main_frame, text="参考Excel文件:").grid(row=0, column=0, sticky=tk.W)
        self.excel_path = tk.StringVar()
        ttk.Entry(self.main_frame, textvariable=self.excel_path, width=40).grid(row=0, column=1)
        ttk.Button(self.main_frame, text="浏览...", command=self.browse_excel).grid(row=0, column=2)

        # 学生文件夹选择
        ttk.Label(self.main_frame, text="学生文件夹:").grid(row=1, column=0, sticky=tk.W)
        self.student_dir = tk.StringVar()
        ttk.Entry(self.main_frame, textvariable=self.student_dir, width=40).grid(row=1, column=1)
        ttk.Button(self.main_frame, text="浏览...", command=self.browse_student_dir).grid(row=1, column=2)

        # 检查按钮
        ttk.Button(self.main_frame, text="开始检查", command=self.run_check).grid(row=2, column=0, columnspan=3, pady=10)

        # 结果显示区域
        self.result_text = tk.Text(self.main_frame, height=15, width=60)
        self.result_text.grid(row=3, column=0, columnspan=3)

        # 滚动条
        scrollbar = ttk.Scrollbar(self.main_frame, orient=tk.VERTICAL, command=self.result_text.yview)
        scrollbar.grid(row=3, column=3, sticky=(tk.N, tk.S))
        self.result_text['yscrollcommand'] = scrollbar.set

    def browse_excel(self):
        filepath = filedialog.askopenfilename(
            title="选择参考Excel文件",
            filetypes=[("Excel文件", "*.xlsx"), ("所有文件", "*.*")]
        )
        if filepath:
            self.excel_path.set(filepath)

    def browse_student_dir(self):
        dirpath = filedialog.askdirectory(title="选择学生文件夹")
        if dirpath:
            self.student_dir.set(dirpath)

    def run_check(self):
        if not self.excel_path.get() or not self.student_dir.get():
            messagebox.showerror("错误", "请先选择参考Excel文件和学生文件夹")
            return

        try:
            # 加载参考数据
            self.ref_data = load_reference_data(self.excel_path.get())

            # 创建结果Excel
            result_wb = Workbook()
            result_ws = result_wb.active
            result_ws.append(['学生姓名', '学号', '异常问题'])

            # 清空结果显示区域
            self.result_text.delete(1.0, tk.END)

            # 检查每个学生
            for student_folder in os.listdir(self.student_dir.get()):
                folder_path = os.path.join(self.student_dir.get(), student_folder)
                if os.path.isdir(folder_path):
                    issues, student_id = check_document_consistency(folder_path, self.ref_data)

                    # 显示结果
                    if issues:
                        self.result_text.insert(tk.END, f"\n学生 {student_id} 发现问题:\n")
                        for issue in issues:
                            self.result_text.insert(tk.END, f"  - {issue}\n")

                    # 写入结果Excel
                    if issues:
                        try:
                            first_doc = next(f for f in os.listdir(folder_path) if f.endswith('.docx'))
                            doc = Document(os.path.join(folder_path, first_doc))
                            info = extract_student_info(doc)
                        except StopIteration:
                            self.result_text.insert(tk.END, f"\n警告: 文件夹 {folder_path} 中没有找到.docx文件\n")
                            continue

                        def format_list(value):
                            if isinstance(value, list):
                                return ', '.join(str(v) for v in value) if value else ''
                            return str(value) if value else ''

                        name_str = format_list(info['id'])
                        advisor_str = format_list(info['advisor'])
                        title_str = format_list(info['title'])

                        for issue in issues:
                            result_ws.append([name_str, student_id, issue])

            # 保存结果
            result_path = os.path.join(os.path.dirname(self.student_dir.get()), "检查结果.xlsx")
            result_wb.save(result_path)
            self.result_text.insert(tk.END, f"\n检查完成，结果已保存到: {result_path}")

        except Exception as e:
            messagebox.showerror("错误", f"检查过程中发生错误:\n{str(e)}")

def main():
    # 加载参考数据
    ref_data = load_reference_data("2023级产业学院毕业设计检查记录表.xlsx")
    # 创建结果Excel
    result_wb = Workbook()
    result_ws = result_wb.active
    result_ws.append(['学生姓名', '学号', '异常问题'])

    # 检查每个学生
    for student_folder in os.listdir("students"):
        folder_path = os.path.join("students", student_folder)
        if os.path.isdir(folder_path):
            issues, student_id = check_document_consistency(folder_path, ref_data)

            # 写入结果
            if issues:
                try:
                    first_doc = next(f for f in os.listdir(folder_path) if f.endswith('.docx'))
                    doc = Document(os.path.join(folder_path, first_doc))
                    info = extract_student_info(doc)
                except StopIteration:
                    print(f"警告: 文件夹 {folder_path} 中没有找到.docx文件")
                    continue

                # 处理所有列表字段为逗号分隔的字符串
                def format_list(value):
                    if isinstance(value, list):
                        return ', '.join(str(v) for v in value) if value else ''
                    return str(value) if value else ''

                name_str = format_list(info['id'])
                advisor_str = format_list(info['advisor'])
                title_str = format_list(info['title'])

                for issue in issues:
                    result_ws.append([name_str, student_id, issue])

    result_wb.save("检查结果.xlsx")
    print("检查完成，结果已保存到'检查结果.xlsx'")

if __name__ == "__main__":
    main()
